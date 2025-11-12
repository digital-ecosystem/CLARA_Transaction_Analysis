"""
Erstellt ein Word-Dokument aus der CSV-Spalten-Erklärung
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading_with_style(doc, text, level):
    """Fügt eine Überschrift mit Stil hinzu"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    return heading

def add_code_block(doc, code_text):
    """Fügt einen Code-Block hinzu"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Hintergrundfarbe für Code
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)
    return p

def create_word_document():
    """Erstellt das Word-Dokument"""
    doc = Document()
    
    # Dokument-Titel
    title = doc.add_heading('CSV-Spalten Erklärung - Analysierte Trades', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Untertitel
    subtitle = doc.add_paragraph('Detaillierte Erklärung aller Spalten in der analysierten CSV-Datei')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.italic = True
    
    # Datum
    date_p = doc.add_paragraph('Stand: 2025-01-12 (nach Trust_Score Entfernung)')
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Leerzeile
    
    # Übersicht
    doc.add_heading('📋 Übersicht', 1)
    doc.add_paragraph(
        'Diese Dokumentation erklärt jede Spalte in der analysierten CSV-Datei '
        '(Analyzed_Trades_*.csv), was die Werte bedeuten und wie sie berechnet werden.'
    )
    
    doc.add_paragraph()  # Leerzeile
    
    # Spalten 1-9: Basis-Informationen
    doc.add_heading('1. Basis-Informationen (Spalten 1-9)', 1)
    
    # Datum
    doc.add_heading('1. Datum (Date)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Das Datum der Transaktion im Format DD.MM.YYYY (z.B. 14.02.2021).')
    p = doc.add_paragraph()
    p.add_run('Quelle: ').bold = True
    p.add_run('Direkt aus der Original-CSV übernommen.')
    p = doc.add_paragraph()
    p.add_run('Berechnung: ').bold = True
    p.add_run('Keine Berechnung, Original-Wert aus der Eingabe-CSV.')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('14.02.2021')
    
    doc.add_paragraph()  # Leerzeile
    
    # Uhrzeit
    doc.add_heading('2. Uhrzeit (Time)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Die Uhrzeit der Transaktion als numerischer Wert (Tagesbruchteil).')
    p = doc.add_paragraph()
    p.add_run('Format: ').bold = True
    p.add_run('0.0 = 00:00:00 (Mitternacht), 0.5 = 12:00:00 (Mittag), 0.630266 = ca. 15:07 Uhr')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Zeitformat (HH:MM:SS)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('0.630266 (entspricht ca. 15:07 Uhr)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Timestamp
    doc.add_heading('3. Timestamp', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Vollständiger Zeitstempel der Transaktion (Datum + Uhrzeit kombiniert).')
    p = doc.add_paragraph()
    p.add_run('Berechnung: ').bold = True
    add_code_block(doc, 'timestamp = Excel_Datum + Uhrzeit\n# Excel_Datum = Anzahl Tage seit 01.01.1900\n# Uhrzeit = Tagesbruchteil (0.0 = 00:00:00, 0.5 = 12:00:00)')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Zahlenformat mit vielen Dezimalstellen')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('44250.630266 (entspricht 14.02.2021 15:07 Uhr)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Kundennummer
    doc.add_heading('4. Kundennummer (Customer Number)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Eindeutige Identifikationsnummer des Kunden.')
    p = doc.add_paragraph()
    p.add_run('Quelle: ').bold = True
    p.add_run('Direkt aus der Original-CSV übernommen.')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('200001')
    
    doc.add_paragraph()  # Leerzeile
    
    # Unique Transaktion ID
    doc.add_heading('5. Unique Transaktion ID', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Eindeutige Identifikationsnummer der Transaktion.')
    p = doc.add_paragraph()
    p.add_run('Format: ').bold = True
    p.add_run('{Kundennummer}-{Timestamp}')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('200001-0,630266203703704')
    
    doc.add_paragraph()  # Leerzeile
    
    # Vollständiger Name
    doc.add_heading('6. Vollständiger Name (Full Name)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Vollständiger Name des Kunden.')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('Bo Lerwagen')
    
    doc.add_paragraph()  # Leerzeile
    
    # Auftragsvolumen
    doc.add_heading('7. Auftragsvolumen (Order Volume)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Der Betrag der Transaktion in EUR.')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Buchhaltungszahlenformat (#,##0.00) mit Tausender-Trennzeichen')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('14000 (14.000,00 EUR)')
    
    doc.add_paragraph()  # Leerzeile
    
    # In/Out
    doc.add_heading('8. In/Out', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Richtung der Transaktion.')
    p = doc.add_paragraph()
    p.add_run('Werte: ').bold = True
    p.add_run('In = Einzahlung/Investment (Geld kommt rein), Out = Auszahlung/Withdrawal (Geld geht raus)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('In')
    
    doc.add_paragraph()  # Leerzeile
    
    # Art
    doc.add_heading('9. Art (Payment Method)', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Zahlungsmethode der Transaktion.')
    p = doc.add_paragraph()
    p.add_run('Werte: ').bold = True
    p.add_run('Bar, SEPA, Kredit, Krypto, etc.')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('SEPA')
    
    doc.add_paragraph()  # Leerzeile
    
    # Analyse-Spalten
    doc.add_heading('2. Analyse-Spalten (Spalten 10-17)', 1)
    
    # Risk_Level
    doc.add_heading('10. Risk_Level ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Das Risiko-Level des Kunden, basierend auf dem Suspicion_Score.')
    p = doc.add_paragraph()
    p.add_run('Werte: ').bold = True
    p.add_run('GREEN = Unauffällig (Suspicion_Score < 150 SP), YELLOW = Leichte Auffälligkeit (150-300 SP), ORANGE = Erhöhtes Risiko (300-500 SP), RED = Hoher Verdacht (≥ 500 SP)')
    add_code_block(doc, 'if suspicion_score < 150:\n    return RiskLevel.GREEN\nelif suspicion_score < 300:\n    return RiskLevel.YELLOW\nelif suspicion_score < 500:\n    return RiskLevel.ORANGE\nelse:\n    return RiskLevel.RED')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Farbcodierung (GREEN=Hellgrün, YELLOW=Gelb, ORANGE=Orange, RED=Rot)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('ORANGE')
    
    doc.add_paragraph()  # Leerzeile
    
    # Suspicion_Score
    doc.add_heading('11. Suspicion_Score ⭐⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Der Suspicion Score (Verdachts-Score) misst das Gesamtrisiko eines Kunden. Höhere Werte = höheres Risiko.')
    p = doc.add_paragraph()
    p.add_run('Berechnung: ').bold = True
    p.add_run('Der Suspicion_Score wird aus mehreren Komponenten berechnet:')
    doc.add_paragraph('1. Module Points (TP/SP System):', style='List Bullet')
    doc.add_paragraph('   - Weight-Analyse (40% Gewicht): Smurfing-Erkennung', style='List Bullet 2')
    doc.add_paragraph('   - Entropy-Analyse (30% Gewicht): Verhaltenskomplexität', style='List Bullet 2')
    doc.add_paragraph('   - Statistics-Analyse (30% Gewicht): Layering, Benford, Velocity, etc.', style='List Bullet 2')
    doc.add_paragraph('2. Gewichtete Summe mit Verstärkungslogik', style='List Bullet')
    doc.add_paragraph('3. Absolute (70%) + Relative (30%) Komponenten', style='List Bullet')
    doc.add_paragraph('4. Nichtlineare Skalierung', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Bereich: ').bold = True
    p.add_run('0.0 bis 1000+ (Suspicion Points)')
    p = doc.add_paragraph()
    p.add_run('Schwellenwerte: ').bold = True
    doc.add_paragraph('< 150 SP: GREEN (Unauffällig)', style='List Bullet')
    doc.add_paragraph('150-300 SP: YELLOW (Leichte Auffälligkeit)', style='List Bullet')
    doc.add_paragraph('300-500 SP: ORANGE (Erhöhtes Risiko)', style='List Bullet')
    doc.add_paragraph('≥ 500 SP: RED (Hoher Verdacht)', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Zahlenformat (0.00) und negativ (z.B. -322.76)')
    p = doc.add_paragraph()
    p.add_run('WICHTIG: ').bold = True
    p.add_run('Der Suspicion_Score wird direkt verwendet (keine Multiplikation oder Division mehr).')
    
    doc.add_paragraph()  # Leerzeile
    
    # Flags
    doc.add_heading('12. Flags ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Spezifische Warnungen und Indikatoren, die für diesen Kunden erkannt wurden.')
    p = doc.add_paragraph()
    p.add_run('Format: ').bold = True
    p.add_run('Mehrere Flags werden mit " | " (Pipe mit Leerzeichen) getrennt.')
    p = doc.add_paragraph()
    p.add_run('Beispiele für Flags: ').bold = True
    doc.add_paragraph('⏱️ HOHE TEMPORALE DICHTE: X.XX Transaktionen/Woche', style='List Bullet')
    doc.add_paragraph('💰 SCHWELLEN-VERMEIDUNG: X.X% der Bar-Investments nah unter 10.000€', style='List Bullet')
    doc.add_paragraph('📍 ENTROPIE-KANALISATION: Extreme Konzentration auf wenige Muster', style='List Bullet')
    doc.add_paragraph('🔴 LAYERING-VERDACHT: Bar → SEPA Muster erkannt', style='List Bullet')
    doc.add_paragraph('👥 PEER-ABWEICHUNG: Untypisch für Kundengruppe', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('⏱️ HOHE TEMPORALE DICHTE: 21.00 Transaktionen/Woche | 📍 ENTROPIE-KANALISATION: Extreme Konzentration auf wenige Muster | 👥 PEER-ABWEICHUNG: Untypisch für Kundengruppe')
    
    doc.add_paragraph()  # Leerzeile
    
    # Threshold_Avoidance_Ratio
    doc.add_heading('13. Threshold_Avoidance_Ratio_% ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Der Anteil (in Prozent) der Bar-Investments, die nah unter der Bar-Grenze (7.000€ - 9.999€) liegen. Ein hoher Wert deutet auf Smurfing hin.')
    add_code_block(doc, '# 1. Finde alle Bar-Investments\nbar_investments = [t for t in transactions\n               if t.payment_method == "Bar" \n               and t.transaction_type == "investment"]\n\n# 2. Finde die, die nah unter der Grenze liegen\nthreshold_avoidance_txns = [t for t in bar_investments\n                            if 7000.0 <= t.transaction_amount < 10000.0]\n\n# 3. Berechne Ratio\nthreshold_avoidance_ratio = len(threshold_avoidance_txns) / len(bar_investments)\nthreshold_avoidance_ratio_percent = threshold_avoidance_ratio * 100')
    p = doc.add_paragraph()
    p.add_run('Bereich: ').bold = True
    p.add_run('0.0% bis 100.0%')
    p = doc.add_paragraph()
    p.add_run('Schwellenwerte: ').bold = True
    doc.add_paragraph('< 30%: Normal', style='List Bullet')
    doc.add_paragraph('≥ 30%: Leicht verdächtig', style='List Bullet')
    doc.add_paragraph('≥ 50%: Verdächtig (starker Smurfing-Indikator)', style='List Bullet')
    doc.add_paragraph('≥ 70%: Sehr verdächtig', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Prozentformat (0.0)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('65.5% (65.5% der Bar-Investments liegen nah unter 10.000€)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Cumulative_Large_Amount
    doc.add_heading('14. Cumulative_Large_Amount ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Die kumulative Summe aller Transaktionen, die nah unter der Bar-Grenze (7.000€ - 9.999€) liegen. Ein hoher Wert deutet auf Smurfing hin.')
    add_code_block(doc, '# 1. Finde alle Bar-Investments nah unter der Grenze\nthreshold_avoidance_txns = [t for t in bar_investments\n                            if 7000.0 <= t.transaction_amount < 10000.0]\n\n# 2. Summiere die Beträge\ncumulative_large_amount = sum(t.transaction_amount \n                              for t in threshold_avoidance_txns)')
    p = doc.add_paragraph()
    p.add_run('Bereich: ').bold = True
    p.add_run('0.00 bis unbegrenzt (in EUR)')
    p = doc.add_paragraph()
    p.add_run('Schwellenwerte: ').bold = True
    doc.add_paragraph('< 30.000€: Normal', style='List Bullet')
    doc.add_paragraph('≥ 30.000€: Leicht verdächtig', style='List Bullet')
    doc.add_paragraph('≥ 50.000€: Verdächtig (starker Smurfing-Indikator)', style='List Bullet')
    doc.add_paragraph('≥ 100.000€: Sehr verdächtig', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Buchhaltungszahlenformat (#,##0.00)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('75.000,00 (75.000€ kumulative Summe nah unter Grenze)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Temporal_Density_Weeks
    doc.add_heading('15. Temporal_Density_Weeks ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Die temporale Dichte der Transaktionen, gemessen als Transaktionen pro Woche. Ein hoher Wert deutet auf verdächtige Aktivität hin.')
    add_code_block(doc, '# 1. Berechne tatsächliche Zeitspanne\ntxns_with_time = [t for t in transactions if t.timestamp]\ntimestamps = [t.timestamp for t in txns_with_time]\nmin_time = min(timestamps)\nmax_time = max(timestamps)\nactual_days = (max_time - min_time).days + 1\nactual_weeks = actual_days / 7.0\n\n# 2. Transaktionen pro Woche\ntemporal_density_weeks = len(txns_with_time) / actual_weeks')
    p = doc.add_paragraph()
    p.add_run('Bereich: ').bold = True
    p.add_run('0.0 bis unbegrenzt (Transaktionen/Woche)')
    p = doc.add_paragraph()
    p.add_run('Schwellenwerte: ').bold = True
    doc.add_paragraph('< 0.25: Normal (weniger als 1 Transaktion/Monat)', style='List Bullet')
    doc.add_paragraph('0.25-0.5: Leicht verdächtig (1-2 Transaktionen/Monat)', style='List Bullet')
    doc.add_paragraph('0.5-1.0: Verdächtig (2-4 Transaktionen/Monat)', style='List Bullet')
    doc.add_paragraph('1.0-2.0: Sehr verdächtig (4-8 Transaktionen/Monat)', style='List Bullet')
    doc.add_paragraph('2.0-5.0: Extrem verdächtig (8-20 Transaktionen/Monat)', style='List Bullet')
    doc.add_paragraph('> 5.0: Sehr extrem verdächtig (mehr als 20 Transaktionen/Monat)', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Zahlenformat (0.00)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('21.0 (21 Transaktionen/Woche = sehr hohe Dichte!)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Layering_Score
    doc.add_heading('16. Layering_Score ⭐⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Der Score für Cash-to-Bank Layering (Geldwäsche-Muster). Misst das Muster: Bar-Einzahlungen → SEPA/Kreditkarte-Auszahlungen.')
    add_code_block(doc, '# 1. Trenne Investments und Auszahlungen\ninvestments = [t for t in transactions if t.transaction_type == "investment"]\nauszahlungen = [t for t in transactions if t.transaction_type == "auszahlung"]\n\n# 2. Finde Bar-Investments und elektronische Auszahlungen\nbar_investments = [t for t in investments if t.payment_method == "Bar"]\nelectronic_withdrawals = [t for t in auszahlungen \n                          if t.payment_method in ["SEPA", "Kreditkarte"]]\n\n# 3. Berechne Verhältnisse\nbar_investment_ratio = len(bar_investments) / len(investments)\nelectronic_withdrawal_ratio = len(electronic_withdrawals) / len(auszahlungen)\n\n# 4. Berechne Volumen-Match und zeitliche Nähe\n# ...\n\n# 5. Kombiniere zu Layering Score\nbase_score = (\n    0.35 * bar_investment_ratio +\n    0.35 * electronic_withdrawal_ratio +\n    0.15 * volume_match_score +\n    0.15 * time_proximity_score\n)\n\n# 6. Absolute Indikatoren (boost)\nif absolute_layering_indicators >= 2:\n    layering_score = min(1.0, base_score + boost)\nelse:\n    layering_score = base_score * 0.3  # Reduktion bei zu wenigen Indikatoren')
    p = doc.add_paragraph()
    p.add_run('Bereich: ').bold = True
    p.add_run('0.0 bis 1.0 (höher = verdächtiger)')
    p = doc.add_paragraph()
    p.add_run('Schwellenwerte: ').bold = True
    doc.add_paragraph('< 0.3: Kein Layering-Verdacht', style='List Bullet')
    doc.add_paragraph('0.3-0.5: Leicht verdächtig', style='List Bullet')
    doc.add_paragraph('0.5-0.7: Verdächtig', style='List Bullet')
    doc.add_paragraph('0.7-0.9: Sehr verdächtig', style='List Bullet')
    doc.add_paragraph('≥ 0.9: Extrem verdächtig (starkes Geldwäsche-Muster)', style='List Bullet')
    p = doc.add_paragraph()
    p.add_run('Excel-Format: ').bold = True
    p.add_run('Zahlenformat (0.00)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('0.85 (85% Layering-Score = sehr verdächtig!)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Entropy_Complex
    doc.add_heading('17. Entropy_Complex ⭐', 2)
    p = doc.add_paragraph()
    p.add_run('Bedeutung: ').bold = True
    p.add_run('Gibt an, ob das Transaktionsverhalten komplex (ungewöhnlich) ist, basierend auf der Entropie-Analyse.')
    add_code_block(doc, '# 1. Berechne Entropie für verschiedene Dimensionen\nentropy_amount = shannon_entropy(amounts)  # Betragsentropie\nentropy_payment_method = shannon_entropy(payment_methods)  # Kanalentropie\nentropy_transaction_type = shannon_entropy(transaction_types)  # Typenentropie\nentropy_time = shannon_entropy(time_patterns)  # Zeitentropie\n\n# 2. Aggregiere Entropie\nentropy_aggregate = (\n    0.30 * entropy_amount +\n    0.30 * entropy_payment_method +\n    0.20 * entropy_transaction_type +\n    0.20 * entropy_time\n)\n\n# 3. Prüfe ob komplex\nis_complex = (entropy_aggregate < 0.3) or (entropy_aggregate > 2.0)\n# - < 0.3: Extreme Konzentration (zu wenig Variation)\n# - > 2.0: Extreme Streuung (zu viel Variation)')
    p = doc.add_paragraph()
    p.add_run('Werte: ').bold = True
    p.add_run('Ja = Entropie ist komplex (entropy_aggregate < 0.3 oder > 2.0), Nein = Entropie ist normal (0.3 ≤ entropy_aggregate ≤ 2.0)')
    p = doc.add_paragraph()
    p.add_run('Beispiel: ').bold = True
    p.add_run('Ja (Extreme Konzentration oder Streuung erkannt)')
    
    doc.add_paragraph()  # Leerzeile
    
    # Zusammenfassung
    doc.add_heading('3. Zusammenfassung', 1)
    
    # Tabelle
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Spalte'
    hdr_cells[1].text = 'Bedeutung'
    hdr_cells[2].text = 'Bereich'
    hdr_cells[3].text = 'Verdächtig wenn'
    
    # Fett für Header
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Zeilen hinzufügen
    rows_data = [
        ('Risk_Level', 'Risiko-Kategorie', 'GREEN/YELLOW/ORANGE/RED', 'ORANGE oder RED'),
        ('Suspicion_Score', 'Gesamt-Verdachts-Score', '0-1000+ SP', '≥ 300 SP'),
        ('Threshold_Avoidance_Ratio_%', 'Anteil nah unter Grenze', '0-100%', '≥ 50%'),
        ('Cumulative_Large_Amount', 'Kumulative Summe', '0-∞ EUR', '≥ 50.000€'),
        ('Temporal_Density_Weeks', 'Transaktionen/Woche', '0-∞', '> 1.0'),
        ('Layering_Score', 'Geldwäsche-Muster', '0.0-1.0', '≥ 0.5'),
        ('Entropy_Complex', 'Komplexes Verhalten', 'Ja/Nein', 'Ja'),
    ]
    
    for row_data in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        row_cells[3].text = row_data[3]
    
    doc.add_paragraph()  # Leerzeile
    
    # Interpretation
    doc.add_heading('4. Interpretation', 1)
    
    doc.add_heading('Normale Kunden', 2)
    doc.add_paragraph('• Risk_Level: GREEN', style='List Bullet')
    doc.add_paragraph('• Suspicion_Score: < 150 SP', style='List Bullet')
    doc.add_paragraph('• Threshold_Avoidance_Ratio: < 30%', style='List Bullet')
    doc.add_paragraph('• Temporal_Density_Weeks: < 0.5', style='List Bullet')
    doc.add_paragraph('• Layering_Score: < 0.3', style='List Bullet')
    doc.add_paragraph('• Entropy_Complex: Nein', style='List Bullet')
    
    doc.add_heading('Verdächtige Kunden', 2)
    doc.add_paragraph('• Risk_Level: YELLOW/ORANGE/RED', style='List Bullet')
    doc.add_paragraph('• Suspicion_Score: ≥ 150 SP', style='List Bullet')
    doc.add_paragraph('• Threshold_Avoidance_Ratio: ≥ 50%', style='List Bullet')
    doc.add_paragraph('• Temporal_Density_Weeks: > 1.0', style='List Bullet')
    doc.add_paragraph('• Layering_Score: ≥ 0.5', style='List Bullet')
    doc.add_paragraph('• Entropy_Complex: Ja', style='List Bullet')
    
    # Speichere Dokument
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_file = f'CSV_SPALTEN_ERKLAERUNG_{timestamp}.docx'
    doc.save(output_file)
    print(f"Word-Dokument erstellt: {output_file}")
    return output_file

if __name__ == '__main__':
    try:
        output_file = create_word_document()
        print(f"\nErfolgreich erstellt: {output_file}")
    except ImportError:
        print("Fehler: python-docx nicht installiert!")
        print("Installiere mit: pip install python-docx")
    except PermissionError:
        print("Fehler: Datei ist geoeffnet oder kein Schreibzugriff!")
        print("Bitte schliessen Sie die Datei und versuchen Sie es erneut.")
    except Exception as e:
        print(f"Fehler: {e}")

