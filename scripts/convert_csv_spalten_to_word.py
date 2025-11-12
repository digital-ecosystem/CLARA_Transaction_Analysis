"""
Konvertiert CSV_SPALTEN_ERKLAERUNG.md zu Word-Format mit professioneller Formatierung
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, code_lines, language=''):
    """Fügt einen formatierten Code-Block hinzu"""
    p = doc.add_paragraph()
    p.style = 'Intense Quote'
    
    # Hintergrundfarbe für Code-Block
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)
    
    # Code-Text hinzufügen
    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)

def parse_markdown_to_word(md_file, docx_file):
    """Konvertiert Markdown zu Word mit professioneller Formatierung"""
    doc = Document()
    
    # Dokument-Stil anpassen
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Titel-Seite
    title = doc.add_heading('CSV-Spalten Erklärung', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Analysierte Trades - CLARA System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # Leerzeile
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Überspringe Titel (wurde bereits hinzugefügt)
        if i < 5 and ('# CSV-Spalten Erklärung' in line or '## 📋 Übersicht' in line):
            i += 1
            continue
        
        # Code-Blöcke
        if line.startswith('```'):
            if in_code_block:
                # Code-Block beenden
                if code_lines:
                    add_code_block(doc, code_lines, code_language)
                code_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Code-Block starten
                code_language = line[3:].strip()
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Überschriften
        if line.startswith('## '):
            heading_text = line[3:].strip()
            # Entferne Emojis
            heading_text = re.sub(r'[⭐📋📊🔍📚]', '', heading_text).strip()
            p = doc.add_heading(heading_text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            heading_text = re.sub(r'[⭐📋📊🔍📚]', '', heading_text).strip()
            p = doc.add_heading(heading_text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            p.runs[0].font.size = Pt(14)
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_text = re.sub(r'[⭐📋📊🔍📚]', '', heading_text).strip()
            p = doc.add_heading(heading_text, level=3)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
            p.runs[0].font.size = Pt(12)
        # Horizontale Linie
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(200, 200, 200)
        # Tabellen
        elif '|' in line and line.strip().startswith('|'):
            # Tabellen-Header erkennen
            if i + 1 < len(lines) and '|' in lines[i + 1] and '---' in lines[i + 1]:
                headers = [cell.strip() for cell in line.split('|')[1:-1]]
                i += 2  # Überspringe Trennlinie
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    # Entferne Emojis und Markdown
                    header_clean = re.sub(r'[⭐📋📊🔍📚]', '', header).strip()
                    header_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', header_clean)
                    header_cells[j].text = header_clean
                    para = header_cells[j].paragraphs[0]
                    para.runs[0].font.bold = True
                    para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    para.runs[0].font.size = Pt(11)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Header-Hintergrundfarbe
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), '4472C4')
                    para._element.get_or_add_pPr().append(shading_elm)
                
                # Tabellen-Zeilen
                while i < len(lines):
                    row_line = lines[i].rstrip()
                    if not row_line.strip().startswith('|') or '---' in row_line:
                        break
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    if len(cells) == len(headers):
                        row = table.add_row()
                        for j, cell in enumerate(cells):
                            # Entferne Markdown-Formatierung
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell)
                            cell_text = re.sub(r'\*(.*?)\*', r'\1', cell_text)
                            cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                            cell_text = re.sub(r'[⭐📋📊🔍📚]', '', cell_text).strip()
                            
                            # Parse für Fettdruck
                            para = row.cells[j].paragraphs[0]
                            para.clear()
                            
                            # Prüfe auf **text** für Fettdruck
                            parts = re.split(r'(\*\*.*?\*\*)', cell)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = para.add_run(part[2:-2])
                                    run.font.bold = True
                                elif part:
                                    para.add_run(part)
                            
                            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    i += 1
                continue
        # Normale Absätze
        elif line.strip():
            # Überspringe leere Zeilen nach Überschriften
            if not line.strip():
                i += 1
                continue
            
            p = doc.add_paragraph()
            
            # Emojis ersetzen
            text = line
            emoji_replacements = {
                '⭐': '★',
                '⚠️': '⚠',
                '📍': '•',
                '🔀': '↔',
                '💳': '💳',
                '📋': '',
                '📊': '',
                '🔍': '',
                '📚': '',
                '🚨': '⚠',
                '💰': '€',
                '💸': '€',
                '🔄': '↻',
                '⏱️': '⏱',
                '🌀': '~',
                '📱': '📱',
                '📉': '↓',
                '👥': '👥'
            }
            for emoji, replacement in emoji_replacements.items():
                text = text.replace(emoji, replacement)
            
            # Spezielle Formatierung für "Bedeutung:", "Quelle:", etc.
            if text.startswith('**Bedeutung:**'):
                run = p.add_run('Bedeutung: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Bedeutung:**', '').strip()
            elif text.startswith('**Quelle:**'):
                run = p.add_run('Quelle: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Quelle:**', '').strip()
            elif text.startswith('**Berechnung:**'):
                run = p.add_run('Berechnung: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Berechnung:**', '').strip()
            elif text.startswith('**Bereich:**'):
                run = p.add_run('Bereich: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Bereich:**', '').strip()
            elif text.startswith('**Schwellenwerte:**'):
                run = p.add_run('Schwellenwerte: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Schwellenwerte:**', '').strip()
            elif text.startswith('**Werte:**'):
                run = p.add_run('Werte: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Werte:**', '').strip()
            elif text.startswith('**WICHTIG:**'):
                run = p.add_run('WICHTIG: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
                text = text.replace('**WICHTIG:**', '').strip()
            elif text.startswith('**Beispiel:**'):
                run = p.add_run('Beispiel: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 102, 0)
                text = text.replace('**Beispiel:**', '').strip()
            elif text.startswith('**Detaillierte Erklärung:**'):
                run = p.add_run('Detaillierte Erklärung: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Detaillierte Erklärung:**', '').strip()
            elif text.startswith('**Komponenten:**'):
                run = p.add_run('Komponenten: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Komponenten:**', '').strip()
            elif text.startswith('**Absolute Indikatoren:**'):
                run = p.add_run('Absolute Indikatoren: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Absolute Indikatoren:**', '').strip()
            elif text.startswith('**Entropie-Bedeutung:**'):
                run = p.add_run('Entropie-Bedeutung: ')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
                text = text.replace('**Entropie-Bedeutung:**', '').strip()
            
            # Markdown-Formatierung parsen
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 100, 0)
                elif part:
                    p.add_run(part)
        
        i += 1
    
    # Speichere Dokument
    doc.save(docx_file)
    print(f"[OK] Word-Dokument erstellt: {docx_file}")

if __name__ == "__main__":
    md_file = "CSV_SPALTEN_ERKLAERUNG.md"
    docx_file = "CSV_SPALTEN_ERKLAERUNG.docx"
    
    try:
        parse_markdown_to_word(md_file, docx_file)
        print(f"\n[OK] Konvertierung erfolgreich!")
        print(f"  Datei: {docx_file}")
    except ImportError:
        print("[FEHLER] python-docx ist nicht installiert")
        print("  Installiere mit: pip install python-docx")
    except Exception as e:
        print(f"[FEHLER] {e}")
        import traceback
        traceback.print_exc()

