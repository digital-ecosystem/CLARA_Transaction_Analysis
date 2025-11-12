"""
Konvertiert SUSPICION_SCORE_BERECHNUNG.md zu Word-Format
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_markdown_to_word(md_file, docx_file):
    """Konvertiert Markdown zu Word"""
    doc = Document()
    
    # Dokument-Stil anpassen
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code-Blöcke
        if line.startswith('```'):
            if in_code_block:
                # Code-Block beenden
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 100, 0)
                    # Hintergrundfarbe für Code
                    p.style = 'Intense Quote'
                code_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Code-Block starten
                code_language = line[3:].strip()
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Überschriften
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        # Horizontale Linie
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('_' * 80).font.color.rgb = RGBColor(200, 200, 200)
        # Tabellen
        elif '|' in line and line.strip().startswith('|'):
            # Tabellen-Header erkennen
            if i + 1 < len(lines) and '|' in lines[i + 1] and '---' in lines[i + 1]:
                headers = [cell.strip() for cell in line.split('|')[1:-1]]
                i += 2  # Überspringe Trennlinie
                
                # Tabelle erstellen
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    header_cells[j].text = header
                    header_cells[j].paragraphs[0].runs[0].font.bold = True
                    header_cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                
                # Tabellen-Zeilen
                while i < len(lines):
                    row_line = lines[i].rstrip()
                    if not row_line.strip().startswith('|') or '---' in row_line:
                        break
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    if len(cells) == len(headers):
                        row = table.add_row()
                        for j, cell in enumerate(cells):
                            # Entferne Markdown-Formatierung
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell)
                            cell_text = re.sub(r'\*(.*?)\*', r'\1', cell_text)
                            row.cells[j].text = cell_text
                            # Fettdruck für **text**
                            if '**' in cell:
                                para = row.cells[j].paragraphs[0]
                                para.clear()
                                parts = re.split(r'(\*\*.*?\*\*)', cell)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        run = para.add_run(part[2:-2])
                                        run.font.bold = True
                                    elif part:
                                        para.add_run(part)
                    i += 1
                continue
        # Normale Absätze
        elif line.strip():
            p = doc.add_paragraph()
            
            # Emojis und Sonderzeichen ersetzen
            text = line
            text = text.replace('⭐', '★')
            text = text.replace('⚠️', '⚠')
            text = text.replace('📍', '•')
            text = text.replace('🔀', '↔')
            text = text.replace('💳', '💳')
            
            # Markdown-Formatierung parsen
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 100, 0)
                elif part:
                    p.add_run(part)
        
        i += 1
    
    # Speichere Dokument
    doc.save(docx_file)
    print(f"[OK] Word-Dokument erstellt: {docx_file}")

if __name__ == "__main__":
    md_file = "SUSPICION_SCORE_BERECHNUNG.md"
    docx_file = "SUSPICION_SCORE_BERECHNUNG.docx"
    
    try:
        parse_markdown_to_word(md_file, docx_file)
        print(f"\n[OK] Konvertierung erfolgreich!")
        print(f"  Datei: {docx_file}")
    except ImportError:
        print("[FEHLER] python-docx ist nicht installiert")
        print("  Installiere mit: pip install python-docx")
    except Exception as e:
        print(f"[FEHLER] {e}")
        import traceback
        traceback.print_exc()
